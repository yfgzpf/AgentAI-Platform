/**
 * 钩子系统管理器
 * 支持流式文档写入、记忆日志、启动脚本等钩子功能
 */

import os
import sys
import json
import asyncio
import subprocess
from pathlib import Path
from typing import Dict, Any, Optional, Callable
from datetime import datetime

try:
    from loguru import logger
except ImportError:
    import logging
    logger = logging.getLogger(__name__)


class HookManager:
    """钩子管理器"""
    
    def __init__(self, hooks_dir: str = None):
        self.hooks_dir = hooks_dir or os.path.join(os.path.expanduser('~'), '.zhiy', 'hooks')
        self.hooks: Dict[str, Callable] = {}
        self._register_builtin_hooks()
        logger.info(f"钩子管理器初始化完成，钩子目录: {self.hooks_dir}")
    
    def _register_builtin_hooks(self):
        """注册内置钩子"""
        self.hooks['stream-doc'] = self._stream_doc_hook
        self.hooks['memory-logger'] = self._memory_logger_hook
        self.hooks['boot-md'] = self._boot_md_hook
        logger.info(f"已注册 {len(self.hooks)} 个内置钩子")
    
    async def run_hook(self, hook_name: str, context: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """
        运行钩子
        
        Args:
            hook_name: 钩子名称
            context: 钩子上下文
        
        Returns:
            钩子执行结果
        """
        hook = self.hooks.get(hook_name)
        if not hook:
            logger.warning(f"钩子 {hook_name} 不存在")
            return None
        
        try:
            result = await hook(context)
            logger.info(f"钩子 {hook_name} 执行成功")
            return result
        except Exception as e:
            logger.error(f"钩子 {hook_name} 执行失败: {str(e)}")
            return {"success": False, "error": str(e)}
    
    async def _stream_doc_hook(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        流式文档写入钩子
        
        当检测到 ##APPEND_DOC:xxx## 标记时，自动将内容追加到文档中
        支持游戏式写入（逐步追加内容）
        """
        result = context.get('result', '')
        session = context.get('session', {})
        session_id = session.get('id', 'default')
        
        # 检测追加标记
        import re
        append_pattern = r'##APPEND_DOC:(.+?)##'
        matches = re.findall(append_pattern, result)
        
        if not matches:
            return {"success": True, "message": "无需追加内容"}
        
        # 创建输出目录
        output_dir = os.path.join(os.path.expanduser('~'), '.zhiy', 'sessions', session_id)
        os.makedirs(output_dir, exist_ok=True)
        doc_path = os.path.join(output_dir, 'output.docx')
        
        results = []
        for text in matches:
            try:
                # 使用 Python 脚本追加内容
                skill_path = os.path.join(os.path.dirname(__file__), '..', '..', '..', 'skills', 'official', 'docx')
                main_script = os.path.join(skill_path, 'main.py')
                
                if os.path.exists(main_script):
                    proc = subprocess.run(
                        ['python', main_script, '--action', 'append', '--file', doc_path, '--text', text],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    results.append({
                        "text": text[:50] + "...",
                        "success": proc.returncode == 0
                    })
                else:
                    # 直接使用 python-docx 追加
                    try:
                        from docx import Document
                        if os.path.exists(doc_path):
                            doc = Document(doc_path)
                        else:
                            doc = Document()
                        
                        doc.add_paragraph(text)
                        doc.save(doc_path)
                        results.append({
                            "text": text[:50] + "...",
                            "success": True
                        })
                    except Exception as e:
                        results.append({
                            "text": text[:50] + "...",
                            "success": False,
                            "error": str(e)
                        })
            except Exception as e:
                results.append({
                    "text": text[:50] + "...",
                    "success": False,
                    "error": str(e)
                })
        
        return {
            "success": all(r["success"] for r in results),
            "message": f"已追加 {len(results)} 段内容到文档",
            "doc_path": doc_path,
            "details": results
        }
    
    async def _memory_logger_hook(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        记忆日志钩子
        
        自动将重要信息记录到记忆文件中
        """
        result = context.get('result', '')
        session = context.get('session', {})
        user_input = context.get('user_input', '')
        
        # 创建记忆目录
        memory_dir = os.path.join(os.path.expanduser('~'), '.zhiy', 'workspace-magic', 'memory')
        os.makedirs(memory_dir, exist_ok=True)
        
        # 写入每日日志
        today = datetime.now().strftime('%Y-%m-%d')
        log_path = os.path.join(memory_dir, f'{today}.md')
        
        timestamp = datetime.now().isoformat()
        log_entry = f"\n## {timestamp}\n### 用户输入\n{user_input}\n\n### AI回复\n{result[:500]}...\n"
        
        try:
            with open(log_path, 'a', encoding='utf-8') as f:
                f.write(log_entry)
            
            return {
                "success": True,
                "message": f"已记录到记忆日志: {log_path}",
                "log_path": log_path
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _boot_md_hook(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        启动脚本钩子
        
        在系统启动时执行自定义脚本
        """
        boot_script = os.path.join(self.hooks_dir, 'boot.md')
        
        if not os.path.exists(boot_script):
            return {
                "success": True,
                "message": "启动脚本不存在，跳过执行"
            }
        
        try:
            with open(boot_script, 'r', encoding='utf-8') as f:
                script_content = f.read()
            
            # 解析并执行脚本中的指令
            lines = script_content.strip().split('\n')
            results = []
            
            for line in lines:
                line = line.strip()
                if line.startswith('#') or not line:
                    continue
                
                # 执行命令
                try:
                    proc = subprocess.run(
                        line,
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=60
                    )
                    results.append({
                        "command": line,
                        "success": proc.returncode == 0,
                        "output": proc.stdout[:200] if proc.stdout else ""
                    })
                except Exception as e:
                    results.append({
                        "command": line,
                        "success": False,
                        "error": str(e)
                    })
            
            return {
                "success": True,
                "message": f"执行了 {len(results)} 条启动指令",
                "results": results
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def register_hook(self, name: str, hook: Callable):
        """注册自定义钩子"""
        self.hooks[name] = hook
        logger.info(f"已注册钩子: {name}")
    
    def unregister_hook(self, name: str):
        """注销钩子"""
        if name in self.hooks:
            del self.hooks[name]
            logger.info(f"已注销钩子: {name}")
    
    def list_hooks(self) -> list:
        """列出所有钩子"""
        return list(self.hooks.keys())


# 全局钩子管理器实例
hook_manager = HookManager()
