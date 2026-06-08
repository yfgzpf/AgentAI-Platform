#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档生成技能
支持创建、编辑 Word 文档
"""

import argparse
import json
import sys
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx 未安装，请运行: pip install python-docx")

def create_document(title: str, content: str, template: str = '标准', output_path: str = None) -> dict:
    """创建 Word 文档"""
    
    if not HAS_DOCX:
        return {
            'status': 'error',
            'message': 'python-docx 库未安装，无法生成 Word 文档'
        }
    
    try:
        doc = Document()
        
        doc.add_heading(title, level=0)
        
        title_para = doc.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. '):
                    doc.add_paragraph(line[3:], style='List Number')
                else:
                    para = doc.add_paragraph(line)
                    para.paragraph_format.first_line_indent = Inches(0.3)
        
        doc.add_paragraph('')
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('由智 Y.Ai 生成')
        
        if not output_path:
            output_dir = os.path.join(os.path.expanduser('~'), 'Documents', 'zhiy-ai')
            os.makedirs(output_dir, exist_ok=True)
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
            output_path = os.path.join(output_dir, f'{safe_title}.docx')
        
        doc.save(output_path)
        
        return {
            'status': 'success',
            'message': f'文档已生成: {output_path}',
            'data': {
                'file_path': output_path,
                'title': title,
                'content_length': len(content)
            }
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'message': f'生成文档失败: {str(e)}'
        }

def append_to_document(file_path: str, content: str) -> dict:
    """追加内容到现有文档"""
    
    if not HAS_DOCX:
        return {
            'status': 'error',
            'message': 'python-docx 库未安装'
        }
    
    try:
        if not os.path.exists(file_path):
            return {
                'status': 'error',
                'message': f'文件不存在: {file_path}'
            }
        
        doc = Document(file_path)
        
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(file_path)
        
        return {
            'status': 'success',
            'message': f'内容已追加到: {file_path}',
            'data': {
                'file_path': file_path,
                'appended_length': len(content)
            }
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'message': f'追加内容失败: {str(e)}'
        }

def main():
    parser = argparse.ArgumentParser(description='Word 文档生成技能')
    parser.add_argument('--title', default='未命名文档', help='文档标题')
    parser.add_argument('--content', default='', help='文档内容')
    parser.add_argument('--template', default='标准', help='模板类型')
    parser.add_argument('--output', default=None, help='输出路径')
    parser.add_argument('--append', action='store_true', help='追加模式')
    parser.add_argument('--file', default=None, help='要追加的文件路径')
    
    args = parser.parse_args()
    
    if args.append and args.file:
        result = append_to_document(args.file, args.content)
    else:
        result = create_document(
            title=args.title,
            content=args.content,
            template=args.template,
            output_path=args.output
        )
    
    print(f"##RESULT## {json.dumps(result, ensure_ascii=False)}")

if __name__ == '__main__':
    main()
