#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档生成技能
支持模板生成、变量替换、流式追加
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx not installed. Run: pip install python-docx")

def create_document(output_path: str, title: str = "未命名文档", content: str = ""):
    """创建新文档"""
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx 未安装"}
    
    doc = Document()
    doc.add_heading(title, 0)
    
    if content:
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
    
    doc.save(output_path)
    return {"success": True, "output": output_path}

def append_to_document(output_path: str, text: str):
    """追加内容到文档"""
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx 未安装"}
    
    if not os.path.exists(output_path):
        doc = Document()
    else:
        doc = Document(output_path)
    
    doc.add_paragraph(text)
    doc.save(output_path)
    return {"success": True, "output": output_path}

def generate_from_template(template_path: str, output_path: str, variables: dict):
    """从模板生成文档"""
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx 未安装"}
    
    if not os.path.exists(template_path):
        return {"success": False, "error": f"模板文件不存在: {template_path}"}
    
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in variables.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    doc.save(output_path)
    return {"success": True, "output": output_path}

def generate_contract(customer_name: str, area: float, style: str, output_path: str):
    """生成装修合同"""
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx 未安装"}
    
    doc = Document()
    
    doc.add_heading('装饰装修合同', 0)
    doc.add_paragraph('')
    
    doc.add_paragraph(f'甲方（客户）：{customer_name}')
    doc.add_paragraph(f'签订日期：{datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph('')
    
    doc.add_heading('一、工程概况', level=1)
    doc.add_paragraph(f'装修面积：{area} 平方米')
    doc.add_paragraph(f'装修风格：{style}')
    doc.add_paragraph('')
    
    doc.add_heading('二、工程内容', level=1)
    doc.add_paragraph('1. 基础装修工程')
    doc.add_paragraph('2. 水电改造工程')
    doc.add_paragraph('3. 墙面处理工程')
    doc.add_paragraph('4. 地面铺设工程')
    doc.add_paragraph('')
    
    doc.add_heading('三、双方责任', level=1)
    doc.add_paragraph('甲方责任：')
    doc.add_paragraph('  1. 按时支付工程款项')
    doc.add_paragraph('  2. 提供施工所需场地')
    doc.add_paragraph('')
    doc.add_paragraph('乙方责任：')
    doc.add_paragraph('  1. 按时保质完成工程')
    doc.add_paragraph('  2. 文明施工，保持现场整洁')
    doc.add_paragraph('')
    
    doc.add_heading('四、付款方式', level=1)
    doc.add_paragraph('1. 合同签订后支付30%预付款')
    doc.add_paragraph('2. 工程过半支付40%进度款')
    doc.add_paragraph('3. 验收合格后支付30%尾款')
    doc.add_paragraph('')
    
    doc.add_paragraph('')
    doc.add_paragraph('甲方签字：________________    日期：________')
    doc.add_paragraph('乙方签字：________________    日期：________')
    
    doc.save(output_path)
    return {"success": True, "output": output_path, "type": "contract"}

def main():
    parser = argparse.ArgumentParser(description='Word文档生成技能')
    parser.add_argument('--action', choices=['create', 'append', 'template', 'contract'], 
                       default='create', help='操作类型')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--template', help='模板文件路径')
    parser.add_argument('--title', default='未命名文档', help='文档标题')
    parser.add_argument('--content', default='', help='文档内容')
    parser.add_argument('--text', default='', help='追加的文本')
    parser.add_argument('--customerName', default='', help='客户姓名')
    parser.add_argument('--area', type=float, default=0, help='面积')
    parser.add_argument('--style', default='', help='风格')
    parser.add_argument('--variables', default='{}', help='JSON格式的变量字典')
    
    args = parser.parse_args()
    
    result = {"success": False, "error": "未知操作"}
    
    try:
        if args.action == 'create':
            result = create_document(args.output, args.title, args.content)
        
        elif args.action == 'append':
            result = append_to_document(args.output, args.text or args.content)
        
        elif args.action == 'template':
            variables = json.loads(args.variables)
            if args.customerName:
                variables['customerName'] = args.customerName
            if args.area:
                variables['area'] = args.area
            if args.style:
                variables['style'] = args.style
            if 'date' not in variables:
                variables['date'] = datetime.now().strftime('%Y-%m-%d')
            
            result = generate_from_template(args.template, args.output, variables)
        
        elif args.action == 'contract':
            result = generate_contract(
                args.customerName or '客户',
                args.area or 100,
                args.style or '现代简约',
                args.output
            )
        
    except Exception as e:
        result = {"success": False, "error": str(e)}
    
    print(f"##RESULT## {json.dumps(result, ensure_ascii=False)}")

if __name__ == '__main__':
    main()
