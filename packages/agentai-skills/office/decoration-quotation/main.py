
#!/usr/bin/env python3
"""
装饰报价技能 - 生成详细的装饰装修报价单
支持材料价格管理和 PDF 导出
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请先安装依赖: pip install python-docx")
    sys.exit(1)


class DecorationQuotation:
    def __init__(self, template_path=None):
        self.template_path = template_path
        self.materials = self._load_default_materials()

    def _load_default_materials(self):
        return {
            "地砖": {"unit": "平方米", "price": 150, "category": "地面材料"},
            "木地板": {"unit": "平方米", "price": 280, "category": "地面材料"},
            "乳胶漆": {"unit": "平方米", "price": 80, "category": "墙面材料"},
            "壁纸": {"unit": "平方米", "price": 120, "category": "墙面材料"},
            "吊顶": {"unit": "平方米", "price": 200, "category": "顶面材料"},
            "橱柜": {"unit": "米", "price": 800, "category": "橱柜"},
            "房门": {"unit": "樘", "price": 1500, "category": "门窗"},
            "卫生间门": {"unit": "樘", "price": 800, "category": "门窗"},
            "厨房门": {"unit": "樘", "price": 600, "category": "门窗"},
            "瓷砖": {"unit": "平方米", "price": 120, "category": "厨房卫生间"},
            "水龙头": {"unit": "个", "price": 300, "category": "厨房卫生间"},
            "马桶": {"unit": "个", "price": 1500, "category": "厨房卫生间"},
            "淋浴花洒": {"unit": "套", "price": 800, "category": "厨房卫生间"},
            "开关插座": {"unit": "个", "price": 30, "category": "水电"},
            "电线": {"unit": "米", "price": 20, "category": "水电"},
            "水管": {"unit": "米", "price": 25, "category": "水电"},
            "人工费": {"unit": "平方米", "price": 300, "category": "人工"},
            "设计费": {"unit": "项", "price": 5000, "category": "其他"},
            "管理费": {"unit": "项", "price": 3000, "category": "其他"},
        }

    def load_materials_from_file(self, filepath):
        if Path(filepath).exists():
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    self.materials = json.load(f)
                return True
            except Exception as e:
                print(f"警告: 无法加载材料文件: {e}", file=sys.stderr)
        return False

    def calculate_item(self, name, quantity, custom_price=None):
        if name not in self.materials:
            return None
        
        material = self.materials[name]
        price = custom_price if custom_price else material["price"]
        total = price * quantity
        
        return {
            "name": name,
            "category": material["category"],
            "unit": material["unit"],
            "price": price,
            "quantity": quantity,
            "total": total
        }

    def generate_quotation(self, customer_name, area, style, items, budget=None):
        quotation = {
            "customer_name": customer_name,
            "area": area,
            "area_unit": "平方米",
            "style": style,
            "date": datetime.now().strftime("%Y年%m月%d日"),
            "items": [],
            "subtotal": 0,
            "tax_rate": 0.06,
            "tax": 0,
            "total": 0,
            "budget": budget
        }

        for item in items:
            calculated = self.calculate_item(
                item["name"],
                item["quantity"],
                item.get("price")
            )
            if calculated:
                quotation["items"].append(calculated)
                quotation["subtotal"] += calculated["total"]

        quotation["tax"] = quotation["subtotal"] * quotation["tax_rate"]
        quotation["total"] = quotation["subtotal"] + quotation["tax"]

        return quotation

    def save_to_docx(self, quotation, output_path):
        doc = Document()

        title = doc.add_heading('装饰装修报价单', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('客户信息', level=1)
        doc.add_paragraph(f'客户姓名: {quotation["customer_name"]}')
        doc.add_paragraph(f'装修面积: {quotation["area"]} {quotation["area_unit"]}')
        doc.add_paragraph(f'户型风格: {quotation["style"]}')
        doc.add_paragraph(f'报价日期: {quotation["date"]}')

        if quotation["budget"]:
            doc.add_paragraph(f'预算范围: {quotation["budget"]}')

        doc.add_heading('报价明细', level=1)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目'
        hdr_cells[1].text = '类别'
        hdr_cells[2].text = '数量'
        hdr_cells[3].text = '单价'
        hdr_cells[4].text = '小计'

        categories = {}
        for item in quotation["items"]:
            cat = item["category"]
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(item)

        for category, items in categories.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f"【{category}】"
            row_cells[1].text = ""
            row_cells[2].text = ""
            row_cells[3].text = ""
            row_cells[4].text = ""
            
            for item in items:
                row_cells = table.add_row().cells
                row_cells[0].text = item["name"]
                row_cells[1].text = item["category"]
                row_cells[2].text = f"{item['quantity']} {item['unit']}"
                row_cells[3].text = f"¥{item['price']:.2f}"
                row_cells[4].text = f"¥{item['total']:.2f}"

        doc.add_heading('费用汇总', level=1)
        doc.add_paragraph(f'小计: ¥{quotation["subtotal"]:.2f}')
        doc.add_paragraph(f'税费 ({quotation["tax_rate"]*100:.0f}%): ¥{quotation["tax"]:.2f}')
        doc.add_paragraph(f'总计: ¥{quotation["total"]:.2f}')

        if quotation["budget"]:
            diff = quotation["total"] - float(quotation["budget"].replace(",", "")) if quotation["budget"] else 0
            if diff &gt; 0:
                doc.add_paragraph(f'⚠️ 超出预算: ¥{diff:.2f}')
            else:
                doc.add_paragraph(f'✅ 预算内: ¥{abs(diff):.2f}')

        doc.add_page_break()
        doc.add_heading('备注', level=1)
        doc.add_paragraph('1. 本报价单仅供参考，实际费用以实际发生为准。')
        doc.add_paragraph('2. 材料价格可能随市场波动，最终价格以实际采购价为准。')
        doc.add_paragraph('3. 施工周期预计为 30-45 天，具体时间根据实际情况确定。')

        try:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"保存失败: {e}", file=sys.stderr)
            return False

    def print_quotation(self, quotation):
        print("=" * 60)
        print("装饰装修报价单".center(60))
        print("=" * 60)
        print(f"客户姓名: {quotation['customer_name']}")
        print(f"装修面积: {quotation['area']} {quotation['area_unit']}")
        print(f"户型风格: {quotation['style']}")
        print(f"报价日期: {quotation['date']}")
        if quotation['budget']:
            print(f"预算范围: {quotation['budget']}")
        print("-" * 60)
        print("\n报价明细:")
        
        categories = {}
        for item in quotation['items']:
            cat = item['category']
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(item)
        
        for category, items in categories.items():
            print(f"\n【{category}】")
            for item in items:
                print(f"  {item['name']:20} {item['quantity']:8} {item['unit']:5} ¥{item['price']:8.2f} = ¥{item['total']:10.2f}")
        
        print("-" * 60)
        print(f"小计: ¥{quotation['subtotal']:.2f}")
        print(f"税费: ¥{quotation['tax']:.2f}")
        print(f"总计: ¥{quotation['total']:.2f}")
        print("=" * 60)


def main():
    parser = argparse.ArgumentParser(description='装饰装修报价生成')
    parser.add_argument('--customer-name', required=True, help='客户姓名')
    parser.add_argument('--area', type=float, required=True, help='装修面积(平方米)')
    parser.add_argument('--style', required=True, help='户型风格')
    parser.add_argument('--budget', help='预算范围')
    parser.add_argument('--items', help='项目清单(JSON格式)')
    parser.add_argument('--items-file', help='项目清单文件(JSON)')
    parser.add_argument('--materials', help='材料价格文件(JSON)')
    parser.add_argument('--output', help='输出Word文件路径')
    parser.add_argument('--template', help='模板文件路径')

    args = parser.parse_args()

    quotation_gen = DecorationQuotation(template_path=args.template)

    if args.materials:
        quotation_gen.load_materials_from_file(args.materials)

    items = []
    if args.items:
        try:
            items = json.loads(args.items)
        except json.JSONDecodeError:
            print("错误: items JSON格式无效", file=sys.stderr)
            sys.exit(1)
    elif args.items_file:
        try:
            with open(args.items_file, 'r', encoding='utf-8') as f:
                items = json.load(f)
        except Exception as e:
            print(f"错误: 无法读取项目文件: {e}", file=sys.stderr)
            sys.exit(1)
    else:
        print("错误: 需要提供项目清单 (--items 或 --items-file)", file=sys.stderr)
        sys.exit(1)

    quotation = quotation_gen.generate_quotation(
        customer_name=args.customer_name,
        area=args.area,
        style=args.style,
        items=items,
        budget=args.budget
    )

    quotation_gen.print_quotation(quotation)

    if args.output:
        if quotation_gen.save_to_docx(quotation, args.output):
            print(f"\n##RESULT## 成功保存到: {args.output}", file=sys.stderr)
        else:
            sys.exit(1)
    else:
        print("\n##RESULT## 成功", file=sys.stderr)


if __name__ == '__main__':
    main()

